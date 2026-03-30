"""
Document processor: extract text from PDF, Excel, DOCX, then chunk for embeddings.
"""
import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from typing import List, Tuple
import io
import re


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    texts = []
    for page in doc:
        texts.append(page.get_text())
    doc.close()
    return "\n".join(texts)


def extract_text_from_excel(file_bytes: bytes) -> str:
    """Extract text from an Excel file — converts all sheets into readable text."""
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    texts = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        texts.append(f"=== Sheet: {sheet_name} ===\n")
        texts.append(df.to_string(index=False))
        texts.append("\n")
    return "\n".join(texts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    doc = Document(io.BytesIO(file_bytes))
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            texts.append(row_text)
    return "\n".join(texts)


def clean_text(text: str) -> str:
    """Clean extracted text by removing excess whitespace and non-printable chars."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x20-\x7E\n]', '', text)
    return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to the correct extractor based on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    
    if ext == "pdf":
        raw = extract_text_from_pdf(file_bytes)
    elif ext in ("xlsx", "xls", "csv"):
        raw = extract_text_from_excel(file_bytes)
    elif ext == "docx":
        raw = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")
    
    return clean_text(raw)


def chunk_text(text: str, chunk_size: int = 800, chunk_overlap: int = 150) -> List[str]:
    """Split text into overlapping chunks for embedding."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    return splitter.split_text(text)


def process_document(file_bytes: bytes, filename: str) -> Tuple[str, List[str]]:
    """
    Full pipeline: extract → clean → chunk.
    Returns (full_text, list_of_chunks).
    """
    full_text = extract_text(file_bytes, filename)
    chunks = chunk_text(full_text)
    return full_text, chunks
